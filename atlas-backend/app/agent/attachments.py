"""Turn user-uploaded files into Gemini input parts.

The chat stream lets the user attach files alongside their question. Two paths,
by type:

  • Images and PDF — sent to Gemini *natively* as ``inline_data`` parts. The
    model sees them (charts, scanned pages, layout), it does not read extracted
    text. No embedding, no retrieval — this is multimodal inference, separate
    from the RAG store in ``app/rag``.
  • Word (.docx) and Excel (.xlsx) — Gemini can't read these binaries, so we
    extract their text/tables here and send that as a plain text part.

Every file is size-capped (``Settings.max_attachment_mb``) before decoding work
happens. A bad file raises :class:`AttachmentError`, which the route turns into a
clean 400 so the frontend can show a precise message.
"""

from __future__ import annotations

import base64
import binascii
import io

from google.genai import types

from app.schema.chat import Attachment
from app.schema.llm_settings import get_settings

# Image types Gemini ingests natively as inline binary (it "sees" them).
_NATIVE_IMAGE_MIMES = {
    "image/png",
    "image/jpeg",
    "image/webp",
    "image/heic",
    "image/heif",
}
_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_XLS_MIME = "application/vnd.ms-excel"  # legacy binary .xls — openpyxl can't read it

# Excel sheets can be enormous; a flat dump blows the context window. Cap the
# rows pulled per sheet and tell the model where it was truncated.
_MAX_SHEET_ROWS = 200


class AttachmentError(ValueError):
    """A bad attachment (oversize, unsupported, undecodable, or empty). Carries a
    stable ``code`` so the frontend can branch on the failure, surfaced as a 400."""

    def __init__(self, message: str, *, code: str) -> None:
        super().__init__(message)
        self.code = code


def build_attachment_parts(attachments: list[Attachment]) -> list[types.Part]:
    """Convert each attachment to a Gemini ``types.Part``. Synchronous and CPU-ish
    (base64 decode + DOCX/XLSX parsing), so the route runs it off the event loop."""
    return [_build_part(att) for att in attachments]


def _decode(att: Attachment) -> bytes:
    try:
        raw = base64.b64decode(att.data, validate=True)
    except (binascii.Error, ValueError) as e:
        raise AttachmentError(
            f"'{att.filename}' is not valid base64 data.", code="bad_encoding"
        ) from e
    cap_mb = get_settings().max_attachment_mb
    if len(raw) > cap_mb * 1024 * 1024:
        raise AttachmentError(
            f"'{att.filename}' is {len(raw) / 1024 / 1024:.1f} MB; the limit is "
            f"{cap_mb} MB per file.",
            code="too_large",
        )
    if not raw:
        raise AttachmentError(f"'{att.filename}' is empty.", code="empty")
    return raw


def _build_part(att: Attachment) -> types.Part:
    raw = _decode(att)
    mime = (att.mime_type or "").lower()
    name = att.filename.lower()

    # Native multimodal: hand the bytes straight to Gemini.
    if mime in _NATIVE_IMAGE_MIMES or mime == _PDF_MIME or name.endswith(".pdf"):
        # Fall back to the PDF mime when the client sent the bytes but no/odd type.
        blob_mime = mime if mime in _NATIVE_IMAGE_MIMES or mime == _PDF_MIME else _PDF_MIME
        return types.Part(inline_data=types.Blob(mime_type=blob_mime, data=raw))

    # Text-extracted: Gemini can't read the Office binary, so we do it here.
    if mime == _DOCX_MIME or name.endswith(".docx"):
        return types.Part(text=_docx_to_text(raw, att.filename))
    if mime == _XLSX_MIME or name.endswith(".xlsx"):
        return types.Part(text=_xlsx_to_text(raw, att.filename))

    if mime == _XLS_MIME or name.endswith(".xls"):
        raise AttachmentError(
            f"'{att.filename}' is a legacy .xls file. Please re-save it as .xlsx.",
            code="unsupported_type",
        )
    raise AttachmentError(
        f"'{att.filename}' ({mime or 'unknown type'}) isn't a supported attachment. "
        "Send an image, PDF, Word (.docx), or Excel (.xlsx) file.",
        code="unsupported_type",
    )


def _docx_to_text(raw: bytes, filename: str) -> str:
    """Paragraphs + table rows from a .docx. Embedded images are dropped (Gemini
    can't see them through this path — send a PDF for image-heavy docs)."""
    import docx

    document = docx.Document(io.BytesIO(raw))
    blocks: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ]
        rows = [r for r in rows if r.strip(" |")]
        if rows:
            blocks.append("\n".join(rows))
    body = "\n\n".join(blocks).strip()
    if not body:
        raise AttachmentError(f"'{filename}' has no readable text.", code="empty")
    return f"Attached Word document '{filename}':\n\n{body}"


def _xlsx_to_text(raw: bytes, filename: str) -> str:
    """Each worksheet as a pipe-delimited table, capped at ``_MAX_SHEET_ROWS``."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    try:
        sheets: list[str] = []
        for ws in wb.worksheets:
            lines: list[str] = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= _MAX_SHEET_ROWS:
                    lines.append(f"... (truncated at {_MAX_SHEET_ROWS} rows)")
                    break
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    lines.append(" | ".join(cells))
            if lines:
                sheets.append(f"### Sheet: {ws.title}\n" + "\n".join(lines))
    finally:
        wb.close()
    body = "\n\n".join(sheets).strip()
    if not body:
        raise AttachmentError(f"'{filename}' has no readable data.", code="empty")
    return f"Attached Excel workbook '{filename}':\n\n{body}"
