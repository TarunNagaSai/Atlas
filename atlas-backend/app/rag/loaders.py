"""Document loaders — PDF (column-aware), DOCX, and plain text."""

from __future__ import annotations

from pathlib import Path

from app.schema.documents import Document

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".pdf", ".docx"}


def load_path(path: str | Path) -> list[Document]:
    p = Path(path)
    if p.is_dir():
        docs: list[Document] = []
        for f in sorted(p.rglob("*")):
            if f.is_file() and f.suffix.lower() in SUPPORTED_SUFFIXES:
                docs.extend(load_file(f))
        return docs
    return load_file(p)


def load_file(path: str | Path) -> list[Document]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(p)
    if suffix == ".docx":
        return _load_docx(p)
    if suffix in TEXT_SUFFIXES:
        return [Document(text=p.read_text(encoding="utf-8", errors="ignore"), source=str(p))]
    raise ValueError(f"Unsupported file type: {p}")


def load_text(text: str, source: str = "inline") -> list[Document]:
    return [Document(text=text, source=source)]


def _extract_page_text(page) -> str:
    """Column-aware extraction: detects two-column layouts by measuring the gap
    between the rightmost word of the left half and the leftmost of the right
    half. When a clear gap exists, each column is cropped and extracted
    independently so text from adjacent columns is never interleaved."""
    words = page.extract_words(x_tolerance=3, y_tolerance=3)
    if not words:
        return ""

    mid = page.width / 2
    left_xs = [w["x1"] for w in words if w["x0"] < mid]
    right_xs = [w["x0"] for w in words if w["x0"] >= mid]

    if left_xs and right_xs and (min(right_xs) - max(left_xs)) >= 10:
        left_text = (
            page.crop((0, 0, max(left_xs) + 5, page.height))
            .extract_text(x_tolerance=3, y_tolerance=3) or ""
        )
        right_text = (
            page.crop((min(right_xs) - 5, 0, page.width, page.height))
            .extract_text(x_tolerance=3, y_tolerance=3) or ""
        )
        return left_text + ("\n\n" + right_text if right_text.strip() else "")

    return page.extract_text(x_tolerance=3, y_tolerance=3) or ""


def _page_pdf_bytes(reader, index: int) -> bytes:
    """Serialize a single page as a standalone one-page PDF (for multimodal
    embedding — the page's text *and* its charts/images in one document)."""
    import io

    import pypdf

    writer = pypdf.PdfWriter()
    writer.add_page(reader.pages[index])
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _load_pdf(path: Path) -> list[Document]:
    import pdfplumber
    import pypdf

    reader = pypdf.PdfReader(str(path))
    docs: list[Document] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            # Keep EVERY page — even image-only/scanned pages with no extractable
            # text. The multimodal embedding captures the page image; the text
            # (when present) still powers lexical search and generation context.
            text = _extract_page_text(page)
            docs.append(
                Document(
                    text=text,
                    source=str(path),
                    metadata={"page": i + 1},
                    embed_pdf=_page_pdf_bytes(reader, i),
                )
            )
    return docs


def _load_docx(path: Path) -> list[Document]:
    import docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text.strip():
        return []
    return [Document(text=text, source=str(path))]
